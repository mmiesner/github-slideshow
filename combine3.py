#!/usr/bin/python3.8
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import os
import sys

#marked this out because I could never figure out the syntax issues to calling it later; didnt work.
#def addstyle(font_name='Times New Roman', font_size=12, before_spacing=0, after_spacing=0, line_spacing=1):
    #paragraph.style = document.styles.add_style('Psych', WD_STYLE_TYPE.PARAGRAPH)
   # font = paragraph.style.font
   # font.name = font_name
    #font.size = Pt(font_size)
    #paragraph_format = paragraph.paragraph_format
    #paragraph_format.space_before = Pt(before_spacing)
   # paragraph_format.space_after = Pt(after_spacing)
   # paragraph.line_spacing = line_spacing

GD = input("What is the proper pronoun for the patient (Mr./Ms.) ?")
firstname = input("What is the patient's first name? ")
lastname = input("What is the patient's last name? ")
DOB = input("What is the patient's date of birth? ")
DOI = input("What is the date of clinical interview date? ")
DOT = input("What is the date of testing? ")
DOR = input("What is the date of the report feedback?")
#BKGRD = input("What is the location of the background file?")
#TSTR =  input("what is the location of the test results file?")
PHYS = input("Who referred the patient?")


document = Document('Psychevaltemplate2.docx')
#addstyle()
#styles = WD_STYLE_TYPE.PARAGRAPH


def find_replace(paragraph_keyword, draft_keyword, paragraph):
    if paragraph_keyword in paragraph.text:
        # print("found")
        paragraph.text = paragraph.text.replace(paragraph_keyword, draft_keyword)


for paragraph in document.paragraphs:
    find_replace("PTFN", firstname, paragraph)
    find_replace("PTLN", lastname, paragraph)
    find_replace("DOBI", DOB, paragraph)
    find_replace("DOI", DOI, paragraph)
    find_replace("DOT", DOT, paragraph)
    find_replace("GD", GD, paragraph)
#for header in document.header:
    #find_replace("PTFN", firstname, paragraph)
    #find_replace("PTLN", lastname, paragraph)


style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']

#addstyle('Psych')
save_filename = "%s, %s Evaluation.docx" % (lastname, firstname)

document.save(save_filename)

document = Document('ReferralLetterTemplate.docx')
for paragraph in document.paragraphs:
    find_replace("PTFN", firstname, paragraph)
    find_replace("PTLN", lastname, paragraph)
    find_replace("DOBI", DOB, paragraph)
    find_replace("DOI", DOI, paragraph)
    find_replace("DOT", DOT, paragraph)
    find_replace("DOR", DOR, paragraph)
    find_replace("PHYS", PHYS, paragraph)
    find_replace("GD", GD, paragraph)

style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']

save_filename = "%s physletter.docx" % lastname

document.save(save_filename)
