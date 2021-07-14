from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import numpy as np

bkgdata = np.loadtxt('background.txt', delimiter='\n', dtype=str)
np.array2string(bkgdata)
bkgdata = [elem.replace("[","").replace("]","") for elem in bkgdata]
bkgdata = [elem.replace("'",'\n').replace("'",'\n') for elem in bkgdata]
print(bkgdata)

document = Document('Psychevaltemplate2.docx')

def find_replace(paragraph_keyword, draft_keyword, paragraph):
    if paragraph_keyword in paragraph.text:
        # print("found")
        paragraph.text = paragraph.text.replace(paragraph_keyword, draft_keyword)


for paragraph in document.paragraphs:
    find_replace("BKGRD", str(bkgdata), paragraph)

style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']


save_filename = "Experimental Full Evaluation.docx"
document.save(save_filename)
exit()

