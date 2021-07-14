#!/usr/bin/python3.8
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import os
import sys
import shutil

GD = input("What is the proper pronoun for the patient (Mr./Ms.) ?")
firstname = input("What is the patient's first name? ")
lastname = input("What is the patient's last name? ")
DOB = input("What is the patient's date of birth? ")
DOI = input("What is the date of clinical interview date? ")
DOT = input("What is the date of testing? ")
DOR = input("What is the date of the report feedback?")
PHYS = input("Who referred the patient?")


document = Document('Psychevaltemplate2.docx')
#addstyle()
#styles = WD_STYLE_TYPE.PARAGRAPH


def find_replace(paragraph_keyword, draft_keyword, paragraph):
    if paragraph_keyword in paragraph.text:
        # print("found")
        paragraph.text = paragraph.text.replace(paragraph_keyword, draft_keyword)


for paragraph in document.paragraphs:
    find_replace("PTFN", firstname, paragraph)
    find_replace("PTLN", lastname, paragraph)
    find_replace("DOBI", DOB, paragraph)
    find_replace("DOI", DOI, paragraph)
    find_replace("DOT", DOT, paragraph)
    find_replace("GD", GD, paragraph)
#for header in document.header:
    #find_replace("PTFN", firstname, paragraph)
    #find_replace("PTLN", lastname, paragraph)


style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']

if sys.platform == 'win32':
   os.chdir('C:\\Users\My PC\Dropbox\written evaluations' )
else: os.chdir('/users/michael/Dropbox/written_evaluations' )

os.mkdir(lastname + ", " + firstname)
os.chdir(lastname + ", " + firstname)

save_filename = "%s Full Evaluation.docx" % lastname
document.save(save_filename)

document = Document(r'C:\Users\My PC\Dropbox\psychpile\ReferralLetterTemplate.docx' )
for paragraph in document.paragraphs:
    find_replace("PTFN", firstname, paragraph)
    find_replace("PTLN", lastname, paragraph)
    find_replace("DOBI", DOB, paragraph)
    find_replace("DOI", DOI, paragraph)
    find_replace("DOT", DOT, paragraph)
    find_replace("DOR", DOR, paragraph)
    find_replace("PHYS", PHYS, paragraph)
    find_replace("GD", GD, paragraph)

style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']

if sys.platform == 'win32':
   os.chdir('C:\\Users\My PC\Dropbox\written evaluations' )
else: os.chdir('/users/michael/Dropbox/written_evaluations' )

os.chdir(lastname + ", " + firstname)

save_filename = "%s physicianletter.docx" % lastname

document.save(save_filename)

document = Document(r'C:\Users\My PC\Dropbox\psychpile\faxsheet.docx' )
for paragraph in document.paragraphs:
    find_replace("PTFN", firstname, paragraph)
    find_replace("PTLN", lastname, paragraph)
    find_replace("DOBI", DOB, paragraph)
    find_replace("DOI", DOI, paragraph)
    find_replace("DOT", DOT, paragraph)
    find_replace("DOR", DOR, paragraph)
    find_replace("PHYS", PHYS, paragraph)
    find_replace("GD", GD, paragraph)

style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']

if sys.platform == 'win32':
   os.chdir('C:\\Users\My PC\Dropbox\written evaluations' )
else: os.chdir('/users/michael/Dropbox/written_evaluations' )

os.chdir(lastname + ", " + firstname)

save_filename = "%s physicianfaxsheet.docx" % lastname

document.save(save_filename)
exit()
