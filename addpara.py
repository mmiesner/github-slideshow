#!/usr/bin/python3.8
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

def get_para_data(output_doc_name, paragraph):
  """
  Write the run to the new file and then set its font, bold, alignment, color etc. data.
  """

  output_para = output_doc_name.add_paragraph()
  for run in paragraph.runs:
    output_run = output_para.add_run(run.text)
    # Run's bold data
    output_run.bold = run.bold
    # Run's italic data
    output_run.italic = run.italic
    # Run's underline data
    output_run.underline = run.underline
    # Run's color data
    #output_run.font.color.rgb = run.font.color.rgb
    # Run's font data
    #output_run.style.name = run.style.name
  # Paragraph's alignment data
  output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment



input_doc = Document('background.docx')
output_doc = Document('Psychevaltemplate3.docx')

# Call the function
get_para_data(output_doc, input_doc.paragraphs[1])
print(input_doc.paragraphs[1])
finalfile = open('Psychevaltemplate3.docx')
finalfile.write(input_doc.paragraphs)
finalfile.close()
# Save the new file

