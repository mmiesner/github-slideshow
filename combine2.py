#!/usr/bin/python3.8
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
def addstyle(font_name = 'Times New Roman', font_size = 12, before_spacing = 0, after_spacing = 0, line_spacing = 1):
    paragraph.style = document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline
    font.color.rgb = color
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before_spacing)
    paragraph_format.space_after = Pt(after_spacing)
    paragraph.line_spacing = line_spacing

addstyle
firstname = input("What is the patient's first name? ")
lastname = input("What is the patient's last name? ")
DOB = input("What is the patient's date of birth? ")
DOI = input("What is the date of clinical interview date? ")
DOT = input("What is the date of testing? ")


document = Document('Psychevaltemplate2.docx')

def find_replace(paragraph_keyword, draft_keyword, paragraph):
    if paragraph_keyword in paragraph.text:
        # print("found")
        paragraph.text = paragraph.text.replace(paragraph_keyword, draft_keyword)


for paragraph in document.paragraphs:
    find_replace("PTFN", firstname, paragraph)
    find_replace("PTLN", lastname, paragraph)
    find_replace("DOB", DOB, paragraph)
    find_replace("DOI", DOI, paragraph)
    find_replace("DOT", DOT, paragraph)


save_filename = "%s, %s Evaluation.docx" % (lastname, firstname)

document.save(save_filename)

document = Document('ReferralLetterTemplate.docx')
addstyle
find_replace("PTFN", firstname, paragraph)
find_replace("PTLN", lastname, paragraph)
find_replace("DOB", DOB, paragraph)
find_replace("DOI", DOI, paragraph)
find_replace("DOT", DOT, paragraph)

save_filename = "%s physletter.docx" % lastname

document.save(save_filename)

