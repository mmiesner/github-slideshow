import PySimpleGUI as sg
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import os
import sys

def docpile(doctype, saveas):
    if doctype = eval
        document =  Document('Psychevaltemplate2.docx')
    if doctype = physletter
        document =  Document(r'C:\Users\My PC\Dropbox\psychpile\ReferralLetterTemplate.docx')
    if doctype = fax
        document = Document(r'C:\Users\My PC\Dropbox\psychpile\faxsheet.docx' )


   if saveas = eval:
       save_filename = "%s Full Evaluation.docx" % lastname
        document.save(save_filename)
   if saveas = physletter: "%s physicianletter.docx" % lastname
         document.save(save_filename)
   if saveas = fax: save_filename = "%s physicianfaxsheet.docx" % lastname
        document.save(save_filename)

    for paragraph in document.paragraphs:
        find_replace("PTFN", firstname, paragraph)
        find_replace("PTLN", lastname, paragraph)
        find_replace("DOBI", DOB, paragraph)
        find_replace("DOI", DOI, paragraph)
        find_replace("DOT", DOT, paragraph)
        find_replace("GD", GD, paragraph)
#for header in document.header:
    #find_replace("PTFN", firstname, paragraph)
    #find_replace("PTLN", lastname, paragraph)


    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    for paragraph in document.paragraphs:
        paragraph.style = document.styles['Normal']


def createf()
    
    if sys.platform == 'win32':
         os.chdir('C:\\Users\My PC\Dropbox\written evaluations' )
    else: os.chdir('/users/michael/Dropbox/written_evaluations' )

    os.chdir(lastname + ", " + firstname)

#end  functions; start iterations

layout = [
    [sg.Text('Lets Compile this PsychFile. Woooo.',justification='center',size=(70,1))],
    [sg.Text('Patient Pronoun', size =(20, 1)), sg.InputText(key='input_GD')],
    [sg.Text('Patient First Name', size =(20, 1)), sg.InputText(key='input_firstname')],
    [sg.Text('Patient Last Name', size =(20, 1)), sg.InputText(key='input_lastname')],
    [sg.Text('Date of Birth', size =(20, 1)), sg.InputText(key='input_DOB')],
    [sg.Text('Clinical Interview Date', size =(20, 1)), sg.InputText(key='input_DOI')],
    [sg.Text('Date of Testing', size =(20, 1)), sg.InputText(key='input_DOT')],
    [sg.Text('Report Feedback Date', size =(20, 1)), sg.InputText(key='input_DOR')],
    [sg.Text('Referrer', size =(20, 1)), sg.InputText(key='input_PHYS')],
    [sg.T("")], [sg.Text("Choose a  background file: "), sg.Input(), sg.FileBrowse(key="input_BKGRD-KEY-")],
    [sg.T("")], [sg.Text("Choose a  testing file: "), sg.Input(), sg.FileBrowse(key="input_TESTING-KEY-")],
    [sg.Button("Submit")]]



window = sg.Window('PsychPile 1.0.0', layout)

while True:
    event, values = window.read()
    #print('event:', event)
    #print('values:', values)
    #print('FolderBrowse:', values['FolderBrowse'])
  #  print('FileBrowse:', values['FileBrowse'])

    if event is None or event == 'Cancel':
        break

    if event == 'Submit':

        GD = (values['input_GD'])
        firstname = (values['input_firstname'])
        lastname = (values['input_lastname'])
        DOB = (values['input_DOB'])
        DOI = (values['input_DOI'])
        DOT = (values['input_DOT'])
        DOR = (values['input_DOR'])
        PHYS = (values['input_PHYS'])
        
        docpile(doctype = eval, saveas = eval)
        createf()

        docpile(doctype = physletter, saveas = physletter)
        createf()

        docpile(doctype = fax, saveas = fax)
        createf()
        


window.close()
exit()
