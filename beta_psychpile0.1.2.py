#!/usr/bin/python3.8
from docx import Document
import os
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

def addstyle(content, font_name = 'Times New Roman', font_size = 12, before_spacing = 0, after_spacing = 0, line_spacing = 1):
    paragraph = document.add_paragraph(str(content))
    paragraph.style = document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline
    font.color.rgb = color
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before_spacing)
    paragraph_format.space_after = Pt(after_spacing)
    paragraph.line_spacing = line_spacing


firstname = input("What is the patient's first name? ")
lastname = input("What is the patient's last name? ")
DOB = input("What is the patient's date of birth? ")
DOI = input("What is the date of clinical interview date? ")
DOT = input("What is the date of testing? ")
DOR = input("What is the date of the report feedback?")
BKGRD = input("What is the location of the background file?")
TSTR =  input("what is the location of the test results file?")
PHYS = input("Who referred the patient?")

print(firstname, lastname, "was born on", DOB, ". The date of testing was", DOT)

document = Document('Psychevaltemplate2.docx')
def find_replace(paragraph_keyword, draft_keyword, paragraph):
    if paragraph_keyword in paragraph.text:
        # print("found")
        paragraph.text = paragraph.text.replace(paragraph_keyword, draft_keyword)
addstyle

for paragraph in document.paragraphs:
    find_replace("PTFN", firstname, paragraph)
    find_replace("PTLN", lastname, paragraph)
    find_replace("DOB", DOB, paragraph)
    find_replace("DOI", DOI, paragraph)
    find_replace("DOT", DOT, paragraph)
    find_replace("DOR", DOR, paragraph)
    #print("text replacement for eval complete")
addstyle
#this is to change PTFN and PTLN in the document header
#for section in document.sections:
    #header = section.header
    #header.paragraphs[1].text  = header.paragraphs[1].text.replace("PTFN", firstname)
    #header.paragraphs[1].text  = header.paragraphs[1].text.replace("PTLN", lastname)
    #print("text replacement for header complete")

#merge background and testing sections
#document = Document(BKGRD)

#for paragraph in document.paragraphs:
#    find_replace("BKGRD", BKGRD, paragraph)
#    find_replace("TSTR", TSTR, paragraph)


filenames = (BKGRD)

var = doc.paragraphs[30]

# Open file3 in write mode 
with open('Psychevaltemplate2.docx', 'w') as outfile:

	# Iterate through list 
	for names in filenames:

		# Open each file in read mode 
		with open(names) as infile:

			# read the data from file1 and 
			# file2 and write it in file3 
			outfile.write(infile.read())

			
filenames = (TSTR)
var = doc.paragraphs[36]

# Open file3 in write mode 
with open('Psychevaltemplate2.docx', 'a') as outfile:

	# Iterate through list 
	for names in filenames:

		# Open each file in read mode 
		with open(names) as infile:

			# read the data from file1 and 
			# file2 and write it in file3 
			outfile.write(infile.read())

save_filename = "%s eval.docx" % lastname
os.path.join('Dropbox', 'written evaluations' '')
save_filename = "%s eval.docx" % lastname
font.name = "Times New Roman"
font.size = Pt(12)
document.save(save_filename)

#fin.close()
fout.close()

#main eval done, now for the physician referral letter
document = Document('ReferralLetterTemplate.docx')

def find_replace(paragraph_keyword, draft_keyword, paragraph):
    if paragraph_keyword in paragraph.text:
        # print("found")
        paragraph.text = paragraph.text.replace(paragraph_keyword, draft_keyword)
        #print("Phys Letter Generated Successfully")

for paragraph in document.paragraphs:
    find_replace("PTFN", firstname, paragraph)
    find_replace("PTLN", lastname, paragraph)
    find_replace("DOB", DOB, paragraph)
    find_replace("DOI", DOI, paragraph)
    find_replace("DOT", DOT, paragraph)
    find_replace("DOR", DOR, paragraph)
    find_replace("PHYS", PHYS, paragraph)
    print(paragraph.text)

save_filename = "%s physletter.docx" % lastname
document.save(save_filename)

#need to create new folders,
# move physletter and generated eval
# to new folder as well

